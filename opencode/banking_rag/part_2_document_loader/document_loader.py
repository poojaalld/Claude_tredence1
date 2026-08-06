"""
Document Loader Module for Banking RAG Assistant
Handles loading and parsing documents from various formats (PDF, DOCX, TXT)
"""

import os
from pathlib import Path
from typing import List, Dict, Optional, Union, Any
from abc import ABC, abstractmethod
from dataclasses import dataclass, asdict
from datetime import datetime
import json

try:
    from langchain.schema import Document
except ImportError:
    from langchain_core.documents import Document


@dataclass
class DocumentMetadata:
    """Metadata for loaded documents"""
    source: str
    file_type: str
    file_size: int
    loaded_at: str
    num_pages: Optional[int] = None
    title: Optional[str] = None
    author: Optional[str] = None
    created_date: Optional[str] = None
    modified_date: Optional[str] = None
    custom_metadata: Optional[Dict[str, Any]] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary"""
        return asdict(self)


class BaseDocumentLoader(ABC):
    """Abstract base class for document loaders"""
    
    def __init__(self, logger=None):
        """
        Initialize the document loader
        
        Args:
            logger: Logger instance for tracking operations
        """
        self.logger = logger
    
    @abstractmethod
    def can_load(self, file_path: Union[str, Path]) -> bool:
        """
        Check if this loader can handle the given file
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if loader can handle the file, False otherwise
        """
        pass
    
    @abstractmethod
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load document from file
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of Document objects
        """
        pass
    
    def _create_metadata(self, file_path: Union[str, Path], **kwargs) -> DocumentMetadata:
        """
        Create metadata for a document
        
        Args:
            file_path: Path to the file
            **kwargs: Additional metadata key-value pairs
            
        Returns:
            DocumentMetadata object
        """
        file_path = Path(file_path)
        stat = file_path.stat()
        
        return DocumentMetadata(
            source=str(file_path),
            file_type=file_path.suffix.lower(),
            file_size=stat.st_size,
            loaded_at=datetime.now().isoformat(),
            modified_date=datetime.fromtimestamp(stat.st_mtime).isoformat(),
            **kwargs
        )
    
    def _log(self, message: str, level: str = "info"):
        """Log a message using the logger"""
        if self.logger:
            getattr(self.logger, level)(message)


class TextDocumentLoader(BaseDocumentLoader):
    """Loader for plain text files"""
    
    def can_load(self, file_path: Union[str, Path]) -> bool:
        """Check if file is a text file"""
        file_path = Path(file_path)
        return file_path.suffix.lower() in ['.txt']
    
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a text file
        
        Args:
            file_path: Path to the text file
            
        Returns:
            List containing a single Document object
        """
        file_path = Path(file_path)
        
        try:
            self._log(f"Loading text file: {file_path}", "debug")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = self._create_metadata(file_path)
            
            doc = Document(
                page_content=content,
                metadata=metadata.to_dict()
            )
            
            self._log(f"Successfully loaded text file: {file_path}", "info")
            return [doc]
            
        except Exception as e:
            self._log(f"Error loading text file {file_path}: {str(e)}", "error")
            raise


class PDFDocumentLoader(BaseDocumentLoader):
    """Loader for PDF files using pypdf"""
    
    def __init__(self, logger=None):
        """Initialize PDF loader"""
        super().__init__(logger)
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are available"""
        try:
            import pypdf
        except ImportError:
            raise ImportError(
                "pypdf is not installed. Install it with: pip install pypdf"
            )
    
    def can_load(self, file_path: Union[str, Path]) -> bool:
        """Check if file is a PDF"""
        file_path = Path(file_path)
        return file_path.suffix.lower() == '.pdf'
    
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a PDF file
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of Document objects, one per page
        """
        from pypdf import PdfReader
        
        file_path = Path(file_path)
        documents = []
        
        try:
            self._log(f"Loading PDF file: {file_path}", "debug")
            
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                metadata = self._create_metadata(file_path, num_pages=num_pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    content = page.extract_text()
                    
                    page_metadata = metadata.to_dict()
                    page_metadata['page_number'] = page_num + 1
                    page_metadata['total_pages'] = num_pages
                    
                    doc = Document(
                        page_content=content,
                        metadata=page_metadata
                    )
                    documents.append(doc)
            
            self._log(f"Successfully loaded PDF file: {file_path} ({num_pages} pages)", "info")
            return documents
            
        except Exception as e:
            self._log(f"Error loading PDF file {file_path}: {str(e)}", "error")
            raise


class DocxDocumentLoader(BaseDocumentLoader):
    """Loader for DOCX (Microsoft Word) files"""
    
    def __init__(self, logger=None):
        """Initialize DOCX loader"""
        super().__init__(logger)
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are available"""
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is not installed. Install it with: pip install python-docx"
            )
    
    def can_load(self, file_path: Union[str, Path]) -> bool:
        """Check if file is a DOCX file"""
        file_path = Path(file_path)
        return file_path.suffix.lower() == '.docx'
    
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List containing a single Document object
        """
        from docx import Document as DocxDocument
        
        file_path = Path(file_path)
        
        try:
            self._log(f"Loading DOCX file: {file_path}", "debug")
            
            doc = DocxDocument(file_path)
            
            # Extract text from all paragraphs
            content = '\n'.join([para.text for para in doc.paragraphs])
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    content += '\n' + ' | '.join(row_data)
            
            metadata = self._create_metadata(file_path)
            
            document = Document(
                page_content=content,
                metadata=metadata.to_dict()
            )
            
            self._log(f"Successfully loaded DOCX file: {file_path}", "info")
            return [document]
            
        except Exception as e:
            self._log(f"Error loading DOCX file {file_path}: {str(e)}", "error")
            raise


class DocumentLoaderFactory:
    """Factory class for creating appropriate document loaders"""
    
    _loaders = [
        TextDocumentLoader,
        PDFDocumentLoader,
        DocxDocumentLoader,
    ]
    
    def __init__(self, logger=None):
        """
        Initialize the factory
        
        Args:
            logger: Logger instance
        """
        self.logger = logger
        self.loaders = []
        
        # Try to load each loader class, skipping those with missing dependencies
        for loader_class in self._loaders:
            try:
                loader_instance = loader_class(logger)
                self.loaders.append(loader_instance)
            except ImportError as e:
                # Log missing dependency but continue
                if logger:
                    logger.warning(f"Loader {loader_class.__name__} skipped: {str(e)}")
                else:
                    print(f"Warning: {loader_class.__name__} skipped: {str(e)}")
    
    def get_loader(self, file_path: Union[str, Path]) -> Optional[BaseDocumentLoader]:
        """
        Get appropriate loader for a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Appropriate loader or None if no loader found
        """
        for loader in self.loaders:
            if loader.can_load(file_path):
                return loader
        return None
    
    def load(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a document using the appropriate loader
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of Document objects
            
        Raises:
            ValueError: If no suitable loader found
        """
        loader = self.get_loader(file_path)
        if not loader:
            raise ValueError(
                f"No loader found for file: {file_path}. "
                f"Supported formats: .txt, .pdf, .docx"
            )
        return loader.load(file_path)


class BulkDocumentLoader:
    """Loader for processing multiple documents"""
    
    def __init__(self, logger=None):
        """
        Initialize bulk loader
        
        Args:
            logger: Logger instance
        """
        self.logger = logger
        self.factory = DocumentLoaderFactory(logger)
        self.loaded_documents = []
        self.failed_files = []
    
    def load_from_directory(
        self, 
        directory_path: Union[str, Path],
        recursive: bool = True,
        extensions: Optional[List[str]] = None
    ) -> List[Document]:
        """
        Load all documents from a directory
        
        Args:
            directory_path: Path to the directory
            recursive: Whether to load recursively from subdirectories
            extensions: List of file extensions to load (e.g., ['.txt', '.pdf'])
            
        Returns:
            List of all loaded Document objects
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise ValueError(f"Directory does not exist: {directory_path}")
        
        all_documents = []
        
        if extensions is None:
            extensions = ['.txt', '.pdf', '.docx']
        
        # Ensure extensions start with dot
        extensions = [ext if ext.startswith('.') else f'.{ext}' for ext in extensions]
        
        self._log(f"Loading documents from: {directory_path}", "info")
        
        # Find all matching files
        if recursive:
            file_pattern = '**/*'
        else:
            file_pattern = '*'
        
        for file_path in directory_path.glob(file_pattern):
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                try:
                    self._log(f"Loading: {file_path}", "debug")
                    docs = self.factory.load(file_path)
                    all_documents.extend(docs)
                    self.loaded_documents.extend(docs)
                except Exception as e:
                    error_msg = f"Failed to load {file_path}: {str(e)}"
                    self._log(error_msg, "error")
                    self.failed_files.append((str(file_path), str(e)))
        
        self._log(
            f"Loaded {len(all_documents)} documents from {len(self.loaded_documents)} files. "
            f"Failed: {len(self.failed_files)}",
            "info"
        )
        
        return all_documents
    
    def load_files(self, file_paths: List[Union[str, Path]]) -> List[Document]:
        """
        Load specific files
        
        Args:
            file_paths: List of file paths to load
            
        Returns:
            List of all loaded Document objects
        """
        all_documents = []
        
        self._log(f"Loading {len(file_paths)} files", "info")
        
        for file_path in file_paths:
            try:
                self._log(f"Loading: {file_path}", "debug")
                docs = self.factory.load(file_path)
                all_documents.extend(docs)
                self.loaded_documents.extend(docs)
            except Exception as e:
                error_msg = f"Failed to load {file_path}: {str(e)}"
                self._log(error_msg, "error")
                self.failed_files.append((str(file_path), str(e)))
        
        self._log(
            f"Loaded {len(all_documents)} documents. Failed: {len(self.failed_files)}",
            "info"
        )
        
        return all_documents
    
    def get_statistics(self) -> Dict[str, Any]:
        """
        Get loading statistics
        
        Returns:
            Dictionary with loading statistics
        """
        total_documents = len(self.loaded_documents)
        total_pages = sum(
            doc.metadata.get('num_pages') or 1
            for doc in self.loaded_documents
        )
        total_size = sum(
            doc.metadata.get('file_size', 0)
            for doc in self.loaded_documents
        )
        
        return {
            'total_documents': total_documents,
            'total_pages': total_pages,
            'total_size_bytes': total_size,
            'failed_count': len(self.failed_files),
            'success_rate': (
                (total_documents - len(self.failed_files)) / total_documents * 100
                if total_documents > 0 else 0
            )
        }
    
    def _log(self, message: str, level: str = "info"):
        """Log a message using the logger"""
        if self.logger:
            getattr(self.logger, level)(message)
